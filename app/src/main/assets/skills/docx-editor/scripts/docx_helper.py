"""Docx-editor helper - programmatic ops on .docx files.

Subcommands:
  inspect <file.docx>                          - per-paragraph + per-table listing
                                                 (idx, style, text preview).
  extract <file.docx>                          - plain text dump.
  set-paragraph <file> --idx N --text "..."    - replace a paragraph's text,
      [--style STYLE] [--out OUT]                preserving the first run's
                                                 font formatting; optionally
                                                 change the paragraph style
                                                 (e.g. 'Heading 1', 'Quote').
  set-cell <file> --table T --row R --col C    - replace a table cell's text,
      --text "..." [--out OUT]                   preserving format.
  new-from-outline <outline.json> --out OUT    - build a docx from a structured
      [--preset NAME]                            outline (heading / paragraph /
                                                 list / table / image /
                                                 page_break + optional header/
                                                 footer + page setup).

Design principles (mirror pptx_helper.py):
  1. NEVER trust paragraph index without `inspect` first.
  2. NEVER use `paragraph.text = "..."` or `cell.text = "..."` -- they wipe
     run-level font/size/bold/color formatting. Helper preserves first run's
     rPr at the XML level.
  3. outline.json gets you 90% of the way; only hand-write python-docx for
     edge cases.

Dependency: python-docx only.
"""
import sys
import os
import json
import argparse
from copy import deepcopy

try:
    if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if sys.stderr.encoding and sys.stderr.encoding.lower() not in ("utf-8", "utf8"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===========================================================================
# Presets (page setup + heading / body fonts)
# ===========================================================================
PRESETS = {
    # name: (page_size, margins_in, body_font, body_size_pt, heading_font, heading_color_hex)
    "business":  ("letter", 1.0, "Calibri",          11, "Calibri",         "1F3864"),
    "academic":  ("letter", 1.0, "Times New Roman",  12, "Times New Roman", "000000"),
    "report":    ("letter", 1.0, "Arial",            11, "Arial",           "C00000"),
    "cjk":       ("a4",     1.0, "Microsoft YaHei",  11, "Microsoft YaHei", "000000"),
}

PAGE_SIZES = {
    "letter": (8.5, 11.0),
    "a4":     (8.27, 11.69),
    "legal":  (8.5, 14.0),
}


# ===========================================================================
# Helpers
# ===========================================================================
def _hex2rgb(h):
    h = h.strip().lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _para_preview(p, max_len=80):
    s = p.text.strip().replace("\n", "\\n")
    if len(s) > max_len:
        s = s[:max_len - 1] + "…"
    return s


# ===========================================================================
# Subcommand: inspect
# ===========================================================================
def cmd_inspect(args):
    doc = Document(args.file)
    # Paragraphs (top-level only, not inside tables)
    print(f"[INSPECT] {args.file}")
    print(f"  paragraphs: {len(doc.paragraphs)}   tables: {len(doc.tables)}   "
          f"sections: {len(doc.sections)}")
    print()
    print(f"─── Paragraphs ───")
    print(f"  {'idx':>3}  {'style':<22}  text")
    print("  " + "-" * 60)
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "?"
        print(f"  {i:>3}  {style:<22}  {_para_preview(p)}")
    print()

    if doc.tables:
        print(f"─── Tables ───")
        for ti, tbl in enumerate(doc.tables):
            nr, nc = len(tbl.rows), len(tbl.columns)
            print(f"  Table {ti}  ({nr} rows x {nc} cols)  "
                  f"style={tbl.style.name if tbl.style else '?'}")
            for r in range(nr):
                for c in range(nc):
                    txt = tbl.cell(r, c).text.strip().replace("\n", "\\n")
                    if len(txt) > 50:
                        txt = txt[:49] + "…"
                    print(f"    ({r},{c}): {txt}")
            print()

    # Sections (page setup)
    print(f"─── Sections ───")
    for si, sec in enumerate(doc.sections):
        pw = sec.page_width.inches if sec.page_width else None
        ph = sec.page_height.inches if sec.page_height else None
        ml = sec.left_margin.inches if sec.left_margin else None
        mr = sec.right_margin.inches if sec.right_margin else None
        mt = sec.top_margin.inches if sec.top_margin else None
        mb = sec.bottom_margin.inches if sec.bottom_margin else None
        print(f"  Section {si}: page {pw}x{ph} in   "
              f"margins L{ml} R{mr} T{mt} B{mb}")
    _next_step([
        "You now have (paragraph_idx, style, text) and (table_idx,row,col,text).",
        f"Replace a paragraph: python docx_helper.py set-paragraph {args.file} --idx N --text \"...\"",
        f"Replace a cell:      python docx_helper.py set-cell {args.file} --table T --row R --col C --text \"...\"",
    ])


# ===========================================================================
# Subcommand: extract
# ===========================================================================
def _iter_block_text(doc):
    """Iterate text in body order: paragraphs and tables interleaved."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            # Find matching paragraph object
            for p in doc.paragraphs:
                if p._p is child:
                    if p.text.strip():
                        yield p.text
                    break
        elif child.tag == qn("w:tbl"):
            for t in doc.tables:
                if t._tbl is child:
                    for row in t.rows:
                        cells = [c.text.strip() for c in row.cells]
                        yield " | ".join(cells)
                    break


def cmd_extract(args):
    doc = Document(args.file)
    for line in _iter_block_text(doc):
        print(line)
    _next_step([
        "Plain text dump above; use for summarization or search.",
        f"To locate paragraph/table indices: python docx_helper.py inspect {args.file}",
    ])


# ===========================================================================
# Preserve-format text replacement
# ===========================================================================
def _preserve_set_paragraph_text(p, new_text):
    """Replace a paragraph's text while keeping the first run's rPr formatting.

    Handles multi-line input by inserting <w:br/> between lines (Word soft
    break, stays in the same paragraph).
    """
    if new_text is None:
        new_text = ""
    lines = new_text.split("\n")

    # Snapshot the first run's rPr xml (if any).
    ref_rPr_xml = None
    if p.runs:
        r0 = p.runs[0]
        rPr = r0._r.find(qn("w:rPr"))
        if rPr is not None:
            ref_rPr_xml = deepcopy(rPr)

    # Remove every <w:r> and <w:hyperlink> child of <w:p> but keep <w:pPr>.
    p_el = p._p
    for child in list(p_el):
        tag = child.tag
        if tag == qn("w:r") or tag == qn("w:hyperlink") or tag == qn("w:fldSimple"):
            p_el.remove(child)

    # Add a single <w:r> with cloned rPr; multi-line via <w:br/>.
    r_el = OxmlElement("w:r")
    if ref_rPr_xml is not None:
        r_el.append(deepcopy(ref_rPr_xml))
    for i, line in enumerate(lines):
        if i > 0:
            br = OxmlElement("w:br")
            r_el.append(br)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r_el.append(t)
    p_el.append(r_el)


# ===========================================================================
# Subcommand: set-paragraph
# ===========================================================================
def cmd_set_paragraph(args):
    # Compat shim: positional idx_pos/text_pos fill legacy attrs.
    if args.idx is None and getattr(args, "idx_pos", None) is not None:
        try:
            args.idx = int(args.idx_pos)
        except (TypeError, ValueError):
            raise ValueError(
                "set-paragraph: IDX must be an integer. Usage:\n"
                "  set-paragraph FILE IDX TEXT"
            )
    if not args.text and getattr(args, "text_pos", None) is not None:
        args.text = args.text_pos
    if args.idx is None or args.text is None:
        raise ValueError(
            "set-paragraph: IDX and TEXT required. Usage:\n"
            "  set-paragraph FILE IDX TEXT [--style \"Heading 1\"]"
        )
    doc = Document(args.file)
    if args.idx < 0 or args.idx >= len(doc.paragraphs):
        raise IndexError(f"paragraph idx {args.idx} out of range "
                         f"(0..{len(doc.paragraphs)-1})")
    p = doc.paragraphs[args.idx]
    if args.style:
        try:
            p.style = doc.styles[args.style]
        except KeyError:
            raise ValueError(f"unknown style '{args.style}'. Common: "
                             f"'Normal', 'Heading 1'..'Heading 9', 'Title', "
                             f"'Subtitle', 'Quote', 'List Bullet', 'List Number'")
    _preserve_set_paragraph_text(p, args.text)
    out = args.out or args.file
    doc.save(out)
    print(f"[OK] set-paragraph idx={args.idx} -> {out}")
    _next_step([
        f"Paragraph {args.idx} updated.",
        f"Verify:   python docx_helper.py inspect {out}",
        "Continue editing other paragraphs with more set-paragraph calls.",
    ])


# ===========================================================================
# Subcommand: set-cell
# ===========================================================================
def cmd_set_cell(args):
    # Compat shim: positional table_pos/row_pos/col_pos/text_pos fill legacy.
    def _to_int(name, raw):
        try:
            return int(raw)
        except (TypeError, ValueError):
            raise ValueError(
                f"set-cell: {name} must be an integer (got {raw!r}). Usage:\n"
                "  set-cell FILE TABLE ROW COL TEXT"
            )
    if args.table is None and getattr(args, "table_pos", None) is not None:
        args.table = _to_int("TABLE", args.table_pos)
    if args.row is None and getattr(args, "row_pos", None) is not None:
        args.row = _to_int("ROW", args.row_pos)
    if args.col is None and getattr(args, "col_pos", None) is not None:
        args.col = _to_int("COL", args.col_pos)
    if not args.text and getattr(args, "text_pos", None) is not None:
        args.text = args.text_pos
    if (args.table is None or args.row is None
            or args.col is None or args.text is None):
        raise ValueError(
            "set-cell: TABLE, ROW, COL, TEXT all required. Usage:\n"
            "  set-cell FILE TABLE ROW COL TEXT"
        )
    doc = Document(args.file)
    if args.table < 0 or args.table >= len(doc.tables):
        raise IndexError(f"table {args.table} out of range "
                         f"(0..{len(doc.tables)-1})")
    tbl = doc.tables[args.table]
    if args.row < 0 or args.row >= len(tbl.rows):
        raise IndexError(f"row {args.row} out of range")
    if args.col < 0 or args.col >= len(tbl.columns):
        raise IndexError(f"col {args.col} out of range")
    cell = tbl.cell(args.row, args.col)
    # A cell can have multiple paragraphs; act on the first, clear the rest.
    if not cell.paragraphs:
        cell.add_paragraph(args.text)
    else:
        _preserve_set_paragraph_text(cell.paragraphs[0], args.text)
        # Drop trailing paragraphs (rarely present in single-line cells).
        for p in cell.paragraphs[1:]:
            p._p.getparent().remove(p._p)
    out = args.out or args.file
    doc.save(out)
    print(f"[OK] set-cell table={args.table} ({args.row},{args.col}) -> {out}")
    _next_step([
        f"Cell ({args.row},{args.col}) of table {args.table} updated.",
        f"Verify:   python docx_helper.py inspect {out}",
    ])


# ===========================================================================
# Subcommand: new-from-outline
# ===========================================================================
def _apply_preset(doc, preset_name):
    if preset_name not in PRESETS:
        raise ValueError(f"unknown preset '{preset_name}'. choices: {list(PRESETS)}")
    page_size, margin_in, body_font, body_pt, head_font, head_color = PRESETS[preset_name]
    pw, ph = PAGE_SIZES[page_size]
    sec = doc.sections[0]
    sec.page_width = Inches(pw); sec.page_height = Inches(ph)
    sec.left_margin = Inches(margin_in); sec.right_margin = Inches(margin_in)
    sec.top_margin = Inches(margin_in); sec.bottom_margin = Inches(margin_in)

    # Set Normal style font
    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal.font.size = Pt(body_pt)
    # CJK font binding (East Asian script)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), body_font)

    # Tweak heading colors (1-3)
    for lvl in range(1, 4):
        try:
            s = doc.styles[f"Heading {lvl}"]
            s.font.name = head_font
            s.font.color.rgb = _hex2rgb(head_color)
            s.font.bold = True
        except KeyError:
            pass


def _add_block(doc, blk):
    t = blk.get("type", "paragraph")
    if t == "heading":
        lvl = int(blk.get("level", 1))
        doc.add_heading(blk.get("text", ""), level=max(0, min(9, lvl)))
    elif t == "paragraph":
        p = doc.add_paragraph(blk.get("text", ""))
        if blk.get("style"):
            try:
                p.style = doc.styles[blk["style"]]
            except KeyError:
                pass
        align = blk.get("align")
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif t == "list":
        items = blk.get("items", [])
        list_style = "List Number" if blk.get("style") == "number" else "List Bullet"
        for it in items:
            try:
                doc.add_paragraph(it, style=list_style)
            except KeyError:
                # Fallback if style not present
                doc.add_paragraph(f"- {it}")
    elif t == "table":
        headers = blk.get("headers", [])
        rows = blk.get("rows", [])
        if not headers or not rows:
            return
        nc = len(headers)
        tbl = doc.add_table(rows=1, cols=nc)
        try:
            tbl.style = blk.get("table_style", "Light Grid Accent 1")
        except KeyError:
            tbl.style = "Table Grid"
        for ci, h in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(h)); r.bold = True
        for row in rows:
            cells = tbl.add_row().cells
            for ci in range(nc):
                v = row[ci] if ci < len(row) else ""
                cells[ci].text = str(v)
    elif t == "image":
        path = blk.get("path")
        if path and os.path.exists(path):
            w_in = float(blk.get("width_in", 5.0))
            doc.add_picture(path, width=Inches(w_in))
        else:
            doc.add_paragraph(f"[image not found: {path}]")
    elif t == "page_break":
        doc.add_page_break()
    else:
        doc.add_paragraph(f"[unknown block type: {t}]")


def _add_page_number_field(paragraph):
    """Append a PAGE field to the paragraph (Word will render the number)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)


VALID_BLOCK_TYPES = {"heading", "paragraph", "list", "table", "image", "page_break"}


def _validate_outline(outline):
    issues = []
    if not isinstance(outline, dict):
        return ["outline must be a JSON object"]
    preset = outline.get("preset", "business")
    if preset not in PRESETS:
        issues.append(f"preset '{preset}' unknown; choices: {list(PRESETS)}")
    blocks = outline.get("blocks", [])
    if not isinstance(blocks, list):
        return issues + ["'blocks' must be an array"]
    for i, b in enumerate(blocks):
        prefix = f"blocks[{i}]"
        if not isinstance(b, dict):
            issues.append(f"{prefix}: must be an object"); continue
        t = b.get("type", "paragraph")
        if t not in VALID_BLOCK_TYPES:
            issues.append(f"{prefix}: unknown type '{t}'; choices: "
                          f"{sorted(VALID_BLOCK_TYPES)}")
            continue
        if t == "heading":
            lvl = b.get("level", 1)
            if not isinstance(lvl, int) or not (0 <= lvl <= 9):
                issues.append(f"{prefix}: heading.level={lvl}, must be int 0..9")
            if not b.get("text"):
                issues.append(f"{prefix}: heading missing 'text'")
        elif t == "paragraph":
            if "text" not in b:
                issues.append(f"{prefix}: paragraph missing 'text'")
        elif t == "list":
            items = b.get("items", [])
            if not items:
                issues.append(f"{prefix}: list.items is empty")
            elif len(items) > 12:
                issues.append(f"{prefix}: list.items={len(items)}, "
                              f"recommend ≤12 (split into multiple lists)")
        elif t == "table":
            headers = b.get("headers", [])
            rows = b.get("rows", [])
            if not headers:
                issues.append(f"{prefix}: table.headers is empty")
            if not rows:
                issues.append(f"{prefix}: table.rows is empty")
            for ri, row in enumerate(rows):
                if isinstance(row, list) and headers and len(row) != len(headers):
                    issues.append(f"{prefix}: table.rows[{ri}] has "
                                  f"{len(row)} cols, headers has "
                                  f"{len(headers)}")
        elif t == "image":
            p = b.get("path")
            if not p:
                issues.append(f"{prefix}: image missing 'path'")
            elif not os.path.exists(p):
                issues.append(f"{prefix}: image not found: {p} "
                              f"(will render placeholder)")

    # Heading level jump check (e.g. h1 → h3 skips h2)
    last_lvl = 0
    for i, b in enumerate(blocks):
        if b.get("type") == "heading":
            lvl = int(b.get("level", 1))
            if lvl > last_lvl + 1 and last_lvl > 0:
                issues.append(f"blocks[{i}]: heading level jumps "
                              f"{last_lvl}→{lvl} (skips levels)")
            last_lvl = lvl
    return issues


# ===========================================================================
# Built-in example outlines (LLM copies + edits, no need to invent structure).
# ===========================================================================
EXAMPLES_DOCX = {
    "report": {
        "_desc": "Business / market report (exec summary + analysis + tables + recommendations).",
        "outline": {
            "title": "Q4 2024 Business Review",
            "author": "Strategy Team",
            "date": "2024-12-31",
            "preset": "business",
            "header": "ACME Corporation · Internal",
            "footer": "Page {page}",
            "blocks": [
                {"type": "heading", "level": 1, "text": "Executive Summary"},
                {"type": "paragraph", "text": "Q4 2024 marked our strongest quarter with 18% YoY revenue growth, NPS climbing 17 points, and successful APAC market entry."},
                {"type": "paragraph", "text": "“Sustained execution on the Q3 roadmap delivered measurable gains across every region.” — CEO", "style": "Quote"},

                {"type": "heading", "level": 1, "text": "Highlights"},
                {"type": "list", "style": "bullet", "items": [
                    "Revenue grew 18% YoY to $5.2M",
                    "Launched 3 major product lines",
                    "Expanded into APAC market",
                    "NPS improved from 45 to 62"
                ]},

                {"type": "heading", "level": 1, "text": "Regional Performance"},
                {"type": "table",
                 "headers": ["Region", "Q4 Revenue", "YoY Growth", "vs Target"],
                 "rows": [
                     ["North America", "$2.8M", "+15%", "Met"],
                     ["Europe", "$1.5M", "+22%", "Exceeded"],
                     ["APAC", "$0.6M", "+85%", "Exceeded"],
                     ["LATAM", "$0.3M", "+5%", "Below"]
                 ]},

                {"type": "page_break"},

                {"type": "heading", "level": 1, "text": "Recommendations"},
                {"type": "heading", "level": 2, "text": "Q1 2025 Priorities"},
                {"type": "list", "style": "number", "items": [
                    "Profitability: target 15% operating margin",
                    "Product: ship v3.0 with AI features",
                    "Hiring: scale engineering to 30",
                    "APAC: localize for JP and KR"
                ]},

                {"type": "heading", "level": 1, "text": "Conclusion"},
                {"type": "paragraph", "text": "Strong fundamentals position us well for FY2025. Continued investment in NPS-driving features and APAC localization will sustain momentum."}
            ]
        }
    },

    "proposal": {
        "_desc": "Project / partnership proposal (problem + solution + scope + timeline + pricing).",
        "outline": {
            "title": "Project Atlas — Enterprise AI Integration Proposal",
            "author": "StartupX Solutions Team",
            "date": "2025-01-15",
            "preset": "business",
            "header": "StartupX · Confidential Proposal",
            "footer": "Page {page}",
            "blocks": [
                {"type": "heading", "level": 1, "text": "Background"},
                {"type": "paragraph", "text": "Client wishes to integrate offline AI capabilities into 50,000 field-deployed devices to reduce cloud dependency and improve response latency."},

                {"type": "heading", "level": 1, "text": "Proposed Solution"},
                {"type": "list", "style": "bullet", "items": [
                    "Deploy our 4B-param edge model to all devices",
                    "Custom domain fine-tuning on client data",
                    "On-device RAG over private knowledge base",
                    "Monitoring dashboard for fleet health"
                ]},

                {"type": "heading", "level": 1, "text": "Scope & Deliverables"},
                {"type": "table",
                 "headers": ["Phase", "Deliverable", "Duration"],
                 "rows": [
                     ["1", "Discovery + data audit", "2 weeks"],
                     ["2", "Fine-tuning + benchmarking", "4 weeks"],
                     ["3", "Pilot deployment (500 devices)", "3 weeks"],
                     ["4", "Full rollout + training", "6 weeks"]
                 ]},

                {"type": "heading", "level": 1, "text": "Timeline"},
                {"type": "paragraph", "text": "Total engagement: 15 weeks from kickoff to full rollout. Pilot results delivered at week 9."},

                {"type": "heading", "level": 1, "text": "Investment"},
                {"type": "table",
                 "headers": ["Item", "Cost"],
                 "rows": [
                     ["Implementation (one-time)", "$240,000"],
                     ["Annual support + updates", "$48,000/yr"],
                     ["Optional: custom dataset curation", "$30,000"]
                 ]},

                {"type": "heading", "level": 1, "text": "Next Steps"},
                {"type": "list", "style": "number", "items": [
                    "Sign mutual NDA",
                    "Kickoff call with technical leads",
                    "Schedule data audit session"
                ]}
            ]
        }
    },

    "minutes": {
        "_desc": "Meeting minutes (attendees + agenda + discussion + action items).",
        "outline": {
            "title": "Q1 Planning Meeting Minutes",
            "author": "Recording Secretary",
            "date": "2025-01-08",
            "preset": "business",
            "blocks": [
                {"type": "heading", "level": 1, "text": "Meeting Details"},
                {"type": "list", "style": "bullet", "items": [
                    "Date: 2025-01-08, 10:00–11:30",
                    "Location: HQ Boardroom + Zoom",
                    "Chair: Jane Doe, CEO",
                    "Recording: John Smith"
                ]},

                {"type": "heading", "level": 1, "text": "Attendees"},
                {"type": "table",
                 "headers": ["Name", "Role", "Status"],
                 "rows": [
                     ["Jane Doe", "CEO", "Present"],
                     ["John Smith", "CTO", "Present"],
                     ["Alex Chen", "VP Product", "Present"],
                     ["Lisa Wang", "VP Sales", "Remote"]
                 ]},

                {"type": "heading", "level": 1, "text": "Agenda"},
                {"type": "list", "style": "number", "items": [
                    "Q4 results review",
                    "Q1 OKRs alignment",
                    "Hiring plan",
                    "AOB"
                ]},

                {"type": "heading", "level": 1, "text": "Discussion"},
                {"type": "heading", "level": 2, "text": "Q4 Results"},
                {"type": "paragraph", "text": "Revenue closed at $5.2M (+18% YoY). All regions met or exceeded targets except LATAM."},

                {"type": "heading", "level": 2, "text": "Q1 OKRs"},
                {"type": "paragraph", "text": "Agreed on three company-level OKRs: 15% op margin, v3.0 launch by March, APAC localization (JP+KR)."},

                {"type": "heading", "level": 1, "text": "Action Items"},
                {"type": "table",
                 "headers": ["Owner", "Action", "Due"],
                 "rows": [
                     ["John", "Finalize v3.0 spec", "2025-01-15"],
                     ["Lisa", "APAC channel partner shortlist", "2025-01-22"],
                     ["Alex", "Hiring brief for 12 roles", "2025-01-12"]
                 ]},

                {"type": "heading", "level": 1, "text": "Next Meeting"},
                {"type": "paragraph", "text": "2025-01-22, 10:00. Same location."}
            ]
        }
    }
}


# ===========================================================================
# _next_step: intentional no-op.
#
# The helper has NO way to know the task's overall goal (is the model going
# to add more blocks? pivot to xlsx? go back to the web for more data?).
# Any "continue with another add-*" / "finish: inspect" suggestion is just
# a guess that pollutes the model's attention.
#
# Scripts should only:
#   * report WHAT HAPPENED    -> [OK] + [STATE]
#   * report errors           -> [ERROR] + targeted recovery hint
#   * provide command menu    -> [HINT] commands list (in main())
# Direction belongs to the model / user prompt, not the helper.
# Kept as a no-op so call sites don't have to be refactored.
# ===========================================================================
def _next_step(lines):  # noqa: ARG001 - intentionally unused
    return


def cmd_examples(args):
    # [DEPRECATED] hint nudging callers toward the new incremental API.
    print("[DEPRECATED] `examples` is the legacy outline-based flow. "
          "Prefer the new incremental API: `catalog` -> `new-doc` -> "
          "`add-*` (one command per block). Run `python docx_helper.py "
          "catalog` to see it.")
    name = args.name
    if not name or name == "list":
        print("Available outline examples (use: examples <name> [--out FILE]):")
        for k, v in EXAMPLES_DOCX.items():
            print(f"  {k:<14}  {v['_desc']}")
        _next_step([
            "Pick a template and dump its outline JSON to a file:",
            "  python docx_helper.py examples <name> --out ${WORKSPACE}/outline.json",
            "Or take the fast path (no outline editing):",
            "  python docx_helper.py create --template <name> --out ${WORKSPACE}/doc.docx",
        ])
        return
    if name not in EXAMPLES_DOCX:
        raise ValueError(f"unknown example '{name}'. choices: "
                         f"{list(EXAMPLES_DOCX)} or 'list'")
    outline_json = json.dumps(EXAMPLES_DOCX[name]["outline"],
                              ensure_ascii=False, indent=2)
    out_path = getattr(args, "out", None)
    if out_path:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(outline_json)
        print(f"[OK] example '{name}' outline written to {out_path}")
        print(f"[WARN] This file is the UNMODIFIED '{name}' template placeholder. "
              f"If you pass it straight to new-from-outline, the build will be REJECTED.")
        _next_step([
            f"REQUIRED next: overwrite {out_path} with your real task content using write_file.",
            "  - Edit 'title', 'author', 'date', and every blocks[i] text fields.",
            "    Keep 'type' unchanged. Tables: edit 'rows'; lists: edit 'items'.",
            f"Then validate: python docx_helper.py new-from-outline {out_path} --dry-run",
            f"Finally build: python docx_helper.py new-from-outline {out_path} --out ${{WORKSPACE}}/doc.docx",
        ])
    else:
        print(outline_json)
        _next_step([
            "Save the JSON above to a file, or re-run with --out FILE to write directly.",
            "Then validate/build with new-from-outline.",
        ])


def _load_json(path):
    with open(path, "r", encoding="utf-8-sig") as f:
        return json.load(f)


def _outline_signature(outline):
    """Stable content hash of an outline, used to detect unmodified templates."""
    return json.dumps(outline, sort_keys=True, ensure_ascii=False)


_EXAMPLE_SIGNATURES = {
    name: _outline_signature(v["outline"]) for name, v in EXAMPLES_DOCX.items()
}


def _detect_unmodified_template(outline):
    """Return example name if outline matches a built-in example byte-for-byte.

    Used to reject `new-from-outline` calls where the LLM forgot to edit the
    example outline produced by `examples NAME --out`, which would otherwise
    build a document containing the template's placeholder copy instead of
    the real task content.
    """
    sig = _outline_signature(outline)
    for name, ref_sig in _EXAMPLE_SIGNATURES.items():
        if sig == ref_sig:
            return name
    return None


def _build_from_outline(outline, out_path, preset_arg=None):
    """Core docx builder; shared by new-from-outline and create.

    Returns (preset_name, block_count).
    """
    preset = preset_arg or outline.get("preset", "business")

    doc = Document()
    _apply_preset(doc, preset)

    # Header / footer (apply to first section)
    sec = doc.sections[0]
    if outline.get("header"):
        sec.header.paragraphs[0].text = outline["header"]
    if outline.get("footer"):
        fp = sec.footer.paragraphs[0]
        ftxt = str(outline["footer"])
        if "{page}" in ftxt:
            before, after = ftxt.split("{page}", 1)
            fp.text = before
            _add_page_number_field(fp)
            fp.add_run(after)
        else:
            fp.text = ftxt
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title block
    if outline.get("title"):
        title_p = doc.add_heading(outline["title"], level=0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if outline.get("author") or outline.get("date"):
        meta = " · ".join(x for x in [outline.get("author"), outline.get("date")] if x)
        sub = doc.add_paragraph(meta)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            sub.style = doc.styles["Subtitle"]
        except KeyError:
            pass

    blocks = outline.get("blocks", [])
    for blk in blocks:
        _add_block(doc, blk)

    doc.save(out_path)
    return preset, len(blocks)


def cmd_new_from_outline(args):
    # [DEPRECATED] hint nudging callers toward the new incremental API.
    print("[DEPRECATED] `new-from-outline` is the legacy single-shot flow. "
          "Prefer the new incremental API: `catalog` -> `new-doc` -> "
          "`add-*` per block. Run `python docx_helper.py catalog` to see it.")
    outline_path = args.outline_pos or args.outline_opt
    if not outline_path:
        raise ValueError("outline path is required "
                         "(pass as positional or --outline PATH)")
    outline = _load_json(outline_path)

    # Guard: refuse to build when the outline is byte-for-byte identical to a
    # built-in example (the LLM forgot to edit it). Otherwise the document
    # would contain the template's placeholder copy, not real task content.
    tmpl = _detect_unmodified_template(outline)
    if tmpl:
        print(f"[ERROR] outline.json is the UNMODIFIED '{tmpl}' example template.")
        print(f"        Building now would produce a generic '{tmpl}' document, NOT your task content.")
        print(f"        You MUST edit the following fields to match the task before re-running:")
        print(f"          - top-level: 'title', 'author', 'date'")
        print(f"          - each blocks[i]: text / rows (tables) / items (lists), keep 'type' unchanged")
        print(f"        Use write_file to overwrite {outline_path} with your customized outline,")
        print(f"        then re-run `new-from-outline`.")
        _next_step([
            f"Overwrite {outline_path} with real task content (keep 'type' unchanged),",
            f"then validate: python docx_helper.py new-from-outline {outline_path} --dry-run",
        ])
        sys.exit(1)

    issues = _validate_outline(outline)
    if issues:
        print("[VALIDATION] issues found:")
        for it in issues:
            print(f"  - {it}")
    else:
        print("[VALIDATION] OK")

    if args.dry_run:
        print("[DRY-RUN] no file written.")
        if issues:
            _next_step([
                "Fix the issues listed above in the outline JSON, then re-run:",
                f"  python docx_helper.py new-from-outline {outline_path} --dry-run",
            ])
            sys.exit(1)
        _next_step([
            "Outline passed validation. To actually build the file, run:",
            f"  python docx_helper.py new-from-outline {outline_path} --out ${{WORKSPACE}}/doc.docx",
        ])
        return

    if not args.out:
        raise ValueError("--out PATH is required unless --dry-run is set")
    preset, n_blocks = _build_from_outline(outline, args.out, args.preset)
    print(f"[OK] new-from-outline -> {args.out} "
          f"(preset={preset}, blocks={n_blocks})")
    _next_step([
        f"Document generated at {args.out}",
        f"Inspect:  python docx_helper.py inspect {args.out}",
        f"Tweak:    python docx_helper.py set-paragraph {args.out} --idx N --text \"...\"",
    ])


def cmd_create(args):
    """Fast path: generate a docx directly from a built-in template."""
    # [DEPRECATED] hint nudging callers toward the new incremental API.
    print("[DEPRECATED] `create --template` produces only the template's "
          "placeholder content (NOT your task content). For real content, "
          "use the new flow: `catalog` -> `new-doc` -> `add-*` per block.")
    if args.outline_file:
        outline = _load_json(args.outline_file)
        src = f"outline file {args.outline_file}"
    else:
        if args.template not in EXAMPLES_DOCX:
            raise ValueError(f"unknown template '{args.template}'. "
                             f"choices: {list(EXAMPLES_DOCX)}")
        outline = deepcopy(EXAMPLES_DOCX[args.template]["outline"])
        src = f"built-in template '{args.template}'"
    issues = _validate_outline(outline)
    if issues:
        print("[VALIDATION] issues:")
        for it in issues:
            print(f"  - {it}")
    else:
        print("[VALIDATION] OK")
    preset, n_blocks = _build_from_outline(outline, args.out, args.preset)
    print(f"[OK] create -> {args.out} (from {src}; "
          f"preset={preset}, blocks={n_blocks})")
    _next_step([
        f"Document generated at {args.out}",
        f"Inspect:  python docx_helper.py inspect {args.out}",
        f"Edit:     python docx_helper.py set-paragraph {args.out} --idx N --text \"...\"",
        "Custom content path: examples <name> --out outline.json "
        "-> edit fields -> new-from-outline <file> --out <doc.docx>",
    ])


# ===========================================================================
# === Catalog / new-doc / add-* incremental API (mirrors pptx_helper) =======
# Per-block additive flow: every `add-*` command appends ONE block (heading /
# paragraph / list / table / image / page-break) and immediately persists the
# .docx, so the model can build a document incrementally without an outline
# JSON. Inputs that exceed the hard caps below are auto-truncated with
# `[WARN]` (these commands NEVER fail on overlong content).
# ===========================================================================

# Hard caps for incremental block API. Numbers picked for visual sanity on a
# typical A4 / Letter page; the model should split content across multiple
# add-* calls when these caps would be hit.
_LIM = {
    "title": 200,           # add-title text
    "subtitle": 300,        # add-title subtitle
    "heading": 200,         # add-heading text
    "paragraph": 8000,      # add-paragraph text (one block)
    "list_item": 200,       # single list item
    "list_items_max": 50,   # max items in one add-list call
    "list_items_min": 1,    # min items (otherwise nothing is added)
    "headers_max": 10,      # max table columns
    "rows_max": 100,        # max table rows per add-table call
    "cell": 500,            # single table cell
}

_META_MARK_DOCX = "docx-helper-doc"


def _meta_default():
    """Default doc-level metadata stored in core_properties.comments."""
    return {"preset": "cjk", "block_count": 0}


def _meta_load(doc):
    """Load doc-level metadata stamped by this helper, with fallbacks."""
    cp = doc.core_properties
    raw = cp.comments or ""
    try:
        m = json.loads(raw) if raw else None
        if isinstance(m, dict) and cp.subject == _META_MARK_DOCX:
            out = _meta_default()
            out.update({k: m[k] for k in out.keys() if k in m})
            out["preset"] = str(out["preset"])
            out["block_count"] = int(out["block_count"])
            return out
    except Exception:
        pass
    fb = _meta_default()
    if cp.subject != _META_MARK_DOCX:
        # Existing doc not produced by this helper -- still allow add-* but
        # seed block_count to current paragraph count so the [OK] message
        # is honest about how many blocks the doc already contains.
        try:
            fb["block_count"] = len(doc.paragraphs)
        except Exception:
            pass
    return fb


def _meta_save(doc, meta):
    """Persist meta back into core_properties (subject + comments)."""
    cp = doc.core_properties
    cp.subject = _META_MARK_DOCX
    cp.comments = json.dumps(meta, ensure_ascii=False)


# --- input clipping (no-fail) ----------------------------------------------
def _clip_str(s, max_chars, *, where=None):
    """Truncate string to max_chars; print [WARN] if trimmed."""
    if s is None:
        return ""
    s = str(s)
    if len(s) <= max_chars:
        return s
    if where:
        print(f"[WARN] {where}: {len(s)} chars > limit {max_chars}, truncated")
    return s[: max_chars - 1] + "\u2026"


def _clip_list(lst, max_n, *, where=None):
    """Drop tail items beyond max_n; print [WARN] if dropped."""
    if not lst:
        return []
    if len(lst) <= max_n:
        return list(lst)
    if where:
        print(f"[WARN] {where}: kept first {max_n} of {len(lst)} items")
    return list(lst[:max_n])


def _split_csv(text, sep=","):
    """Split a comma-separated arg into stripped non-empty tokens."""
    if not text:
        return []
    return [t.strip() for t in str(text).split(sep) if t.strip()]


def _open_doc(path):
    """Load a docx; if missing, instruct caller to run new-doc first."""
    if not os.path.exists(path):
        raise FileNotFoundError(
            "Doc not found: " + str(path) + "\n"
            "Hint: run `docx_helper.py new-doc " + str(path) +
            " --preset cjk` first."
        )
    return Document(path)


def _commit_block(doc, args, kind, meta, *, state_extra=None):
    """Persist doc after adding a block and print [OK] / NEXT_STEP / [STATE]."""
    meta["block_count"] = int(meta.get("block_count", 0)) + 1
    _meta_save(doc, meta)
    doc.save(args.file)
    n = meta["block_count"]
    print(f"[OK] {kind} -> {args.file} "
          f"(blocks={n}, preset={meta['preset']})")
    # A4: tell the model the real document shape so it can decide whether
    # to keep adding or move on (no need to call inspect).
    state_line = f"[STATE] doc='{args.file}' blocks={n}"
    if state_extra:
        state_line += " " + state_extra
    print(state_line)
    _next_step([
        f"Block #{n} added. Continue with another add-* "
        f"(e.g. add-heading, add-paragraph, add-list, add-table)",
        f"or finish: python docx_helper.py inspect {args.file}",
    ])


# ---------------------------------------------------------------------------
# A1 / A2 / A3 helpers (mirror of xlsx_helper -- log2.txt 15:40:59 fix)
# ---------------------------------------------------------------------------
def _expand_packed_rows_docx(rows):
    """A1: positional row tolerance for add-table.

    The model regularly squeezes multiple csv rows into ONE positional
    arg joined by '|' or ';'. Detect and split using the same heuristic
    as xlsx_helper._expand_packed_rows. Returns (expanded, n_split).
    """
    if not rows:
        return list(rows or []), 0
    out = []
    n_split = 0
    for r in rows:
        s = str(r)
        if (("|" in s) or (";" in s)) and ("," in s):
            sep = "|" if "|" in s else ";"
            parts = [p.strip() for p in s.split(sep) if p.strip()]
            if len(parts) >= 2 and all(("," in p) for p in parts):
                out.extend(parts)
                n_split += 1
                continue
        out.append(s)
    return out, n_split


def _warn_packed_rows_docx(n_split):
    if n_split <= 0:
        return
    print(
        f"[WARN][POSITIONAL_ROW_DELIMITER] add-table: detected {n_split} "
        f"positional argument(s) packed with '|' or ';' as row separator. "
        f"Auto-split into multiple rows."
    )
    print(
        "  Canonical form (preferred):\n"
        "    --row \"a,b\" --row \"c,d\"             (one --row per row)\n"
        "    --rows \"a,b|c,d\"                     (alias, '|' splits rows)\n"
        "  Positional form (one row per arg):\n"
        "    add-table FILE \"hdr1,hdr2\" \"a,b\" \"c,d\"  (each arg is ONE row)"
    )


def _warn_col_mismatch_docx(*, expected, got, where):
    if expected <= 0 or got == expected:
        return
    print(
        f"[WARN][COL_MISMATCH] add-table {where}: expected {expected} "
        f"column(s) (per --headers), got {got}."
    )
    if got < expected:
        print(
            f"  Missing {expected - got} cell(s) -- written as blanks. "
            f"Pad your CSV: \"v1,v2,...,v{expected}\""
        )
    else:
        print(
            f"  Extra {got - expected} cell(s) -- truncated. Drop the "
            f"trailing values or fix --headers to {got} columns."
        )


def _row_signature_docx(values):
    return tuple(("" if v is None else str(v).strip()) for v in values)


def _warn_duplicate_row_docx(*, prev_idx, cur_idx, signature):
    preview = ",".join(list(signature)[:4])
    if len(signature) > 4:
        preview += ",..."
    print(
        f"[WARN][DUPLICATE_ROW] add-table: row {cur_idx} is identical to "
        f"row {prev_idx} ({preview}). Wrote anyway."
    )


# --- new-doc ---------------------------------------------------------------
def cmd_new_doc(args):
    """Initialize an empty .docx with a locked preset."""
    preset = (args.preset or "cjk").strip().lower()
    if preset not in PRESETS:
        print(f"[WARN] preset '{preset}' unknown; falling back to 'cjk' "
              f"(valid: {list(PRESETS)})")
        preset = "cjk"
    doc = Document()
    _apply_preset(doc, preset)
    _meta_save(doc, {"preset": preset, "block_count": 0})
    doc.save(args.file)
    print(f"[OK] new-doc -> {args.file} (preset={preset}, blocks=0)")
    _next_step([
        "Add blocks one at a time (each command appends one block):",
        f"  python docx_helper.py add-title {args.file} "
        f"--text \"Your Title\" --subtitle \"Your Subtitle\"",
        f"  python docx_helper.py add-heading {args.file} "
        f"--level 1 --text \"Section\"",
        f"  python docx_helper.py add-paragraph {args.file} "
        f"--text \"Body text...\"",
        f"  python docx_helper.py add-list {args.file} "
        f"--type bullet --item \"point 1\" --item \"point 2\"",
        f"  python docx_helper.py add-table {args.file} "
        f"--headers \"a,b,c\" --row \"1,2,3\"",
    ])


# --- add-title -------------------------------------------------------------
def cmd_add_title(args):
    """Append a centered title (+ optional subtitle). Use ONCE at start."""
    doc = _open_doc(args.file)
    meta = _meta_load(doc)
    # Resolve title text from --text first, then positional fallback.
    title_raw = args.text or getattr(args, "pos_text", None)
    if not (title_raw and str(title_raw).strip()):
        sys.stderr.write("[ERROR] add-title: missing text. "
                         "Use --text \"...\" or positional: "
                         "add-title FILE.docx \"Title\" [\"Subtitle\"]\n")
        _print_commands_menu(stream=sys.stderr)
        sys.exit(2)
    sub_raw = args.subtitle or getattr(args, "pos_subtitle", None)
    title = _clip_str(title_raw, _LIM["title"], where="title text")
    subtitle = _clip_str(sub_raw, _LIM["subtitle"], where="subtitle")
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        sp = doc.add_paragraph(subtitle)
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            sp.style = doc.styles["Subtitle"]
        except KeyError:
            pass
    _commit_block(doc, args, "add-title", meta)


# --- add-heading -----------------------------------------------------------
def cmd_add_heading(args):
    """Append a heading at level 1..9 (use level=1 for top-level sections).

    LLM-friendly invocation:
        add-heading FILE.docx --text "Section" --level 1   (canonical)
        add-heading FILE.docx "Section" 1                  (positional)
        add-heading FILE.docx 1 "Section"                  (swapped order ok)
    """
    doc = _open_doc(args.file)
    meta = _meta_load(doc)
    # Sniff level vs text from --text/--level + positional fallbacks.
    text_raw = args.text
    lvl_raw = args.level
    for cand in (getattr(args, "pos1", None), getattr(args, "pos2", None)):
        if cand is None:
            continue
        # Pure-int candidate -> treat as level (only if level is still missing)
        try:
            cand_int = int(str(cand))
            if lvl_raw is None and 0 <= cand_int <= 9:
                lvl_raw = cand_int
                continue
        except (TypeError, ValueError):
            pass
        # Otherwise -> text fallback (only if text is still missing)
        if not (text_raw and str(text_raw).strip()):
            text_raw = cand
    if not (text_raw and str(text_raw).strip()):
        sys.stderr.write("[ERROR] add-heading: missing text. "
                         "Use --text \"...\" or positional: "
                         "add-heading FILE.docx \"Section\" [LEVEL]\n")
        _print_commands_menu(stream=sys.stderr)
        sys.exit(2)
    try:
        lvl = int(lvl_raw) if lvl_raw is not None else 1
    except (TypeError, ValueError):
        lvl = 1
    lvl = max(1, min(9, lvl))
    text = _clip_str(text_raw, _LIM["heading"], where="heading text")
    doc.add_heading(text, level=lvl)
    _commit_block(doc, args, "add-heading", meta)


# --- add-paragraph ---------------------------------------------------------
def cmd_add_paragraph(args):
    """Append a paragraph (optional --style and --align)."""
    doc = _open_doc(args.file)
    meta = _meta_load(doc)
    text_raw = args.text or getattr(args, "pos_text", None)
    if not (text_raw and str(text_raw).strip()):
        sys.stderr.write("[ERROR] add-paragraph: missing text. "
                         "Use --text \"...\" or positional: "
                         "add-paragraph FILE.docx \"Body...\"\n")
        _print_commands_menu(stream=sys.stderr)
        sys.exit(2)
    text = _clip_str(text_raw, _LIM["paragraph"], where="paragraph text")
    p = doc.add_paragraph(text)
    style = (args.style or "").strip()
    if style:
        try:
            p.style = doc.styles[style]
        except KeyError:
            print(f"[WARN] style '{style}' not in doc.styles; using Normal")
    align = (args.align or "").lower()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _commit_block(doc, args, "add-paragraph", meta)


# --- add-list --------------------------------------------------------------
# Synonym map for --type (the LLM frequently writes 'unordered'/'ol'/'numbered').
_LIST_TYPE_ALIAS = {
    "bullet": "bullet", "bullets": "bullet",
    "unordered": "bullet", "ul": "bullet", "disc": "bullet",
    "number": "number", "numbered": "number",
    "ordered": "number", "ol": "number",
}


def cmd_add_list(args):
    """Append a bullet (default) or numbered list. 1..50 items."""
    doc = _open_doc(args.file)
    meta = _meta_load(doc)
    # Merge --item (canonical) with positional pos_items (LLM fallback).
    items_raw = list(args.item or []) + list(getattr(args, "pos_items", []) or [])
    items_raw = _clip_list(items_raw, _LIM["list_items_max"],
                           where="list items")
    items = [_clip_str(it, _LIM["list_item"], where=f"list[{i}]")
             for i, it in enumerate(items_raw) if str(it).strip()]
    if len(items) < _LIM["list_items_min"]:
        print("[WARN] add-list: need at least 1 non-empty item "
              "(via --item or positional); nothing added")
        return
    type_raw = (args.type or "bullet").strip().lower()
    type_norm = _LIST_TYPE_ALIAS.get(type_raw)
    if type_norm is None:
        print(f"[WARN] add-list: unknown --type '{args.type}', "
              f"defaulting to 'bullet' (accepted: bullet/number, "
              f"aliases: unordered/ul -> bullet, ordered/ol/numbered -> number)")
        type_norm = "bullet"
    list_style = "List Number" if type_norm == "number" else "List Bullet"
    for it in items:
        try:
            doc.add_paragraph(it, style=list_style)
        except KeyError:
            doc.add_paragraph(f"- {it}")
    _commit_block(doc, args, "add-list", meta)


# --- add-table -------------------------------------------------------------
def cmd_add_table(args):
    """Append a headed table. Headers/rows are comma-separated.

    LLM-friendly invocation forms accepted:
        add-table FILE.docx --headers "a,b" --row "1,2" --row "3,4"  (canonical)
        add-table FILE.docx --headers "a,b" --rows "1,2|3,4"        (alias --rows splits on '|')
        add-table FILE.docx "a,b" "1,2" "3,4"                       (positional fallback)
    """
    doc = _open_doc(args.file)
    meta = _meta_load(doc)
    # Merge keyword args with positional fallback.
    pos_args = list(getattr(args, "pos_args", []) or [])
    headers_str = args.headers
    if headers_str is None and pos_args:
        headers_str = pos_args.pop(0)
    if not (headers_str and str(headers_str).strip()):
        sys.stderr.write("[ERROR] add-table: missing headers. Use "
                         "--headers \"a,b\" or positional: "
                         "add-table FILE.docx \"a,b\" \"1,2\" ...\n")
        _print_commands_menu(stream=sys.stderr)
        sys.exit(2)
    headers = _split_csv(headers_str)
    headers = _clip_list(headers, _LIM["headers_max"], where="table headers")
    headers = [_clip_str(h, _LIM["cell"], where=f"header[{i}]")
               for i, h in enumerate(headers)]
    # Aggregate rows from --row (canonical), --rows (alias, '|'-split), and
    # any leftover positional args. Empty values dropped.
    rows_raw = list(args.row or [])
    for rs in (getattr(args, "rows", []) or []):
        rows_raw.extend([s for s in str(rs).split("|") if s.strip()])
    # [A1] Apply packed-row tolerance to leftover positional args BEFORE
    # appending them. log2.txt 15:40:59: model packed 8 DRAM rows into a
    # single positional containing '|' separators -> we now split it.
    pos_args_clean = [s for s in pos_args if str(s).strip()]
    pos_args_expanded, n_split = _expand_packed_rows_docx(pos_args_clean)
    _warn_packed_rows_docx(n_split)
    rows_raw.extend(pos_args_expanded)
    rows_raw = _clip_list(rows_raw, _LIM["rows_max"], where="table rows")
    rows = []
    raw_signatures = []  # for A3
    for ri, row_str in enumerate(rows_raw):
        cells = _split_csv(row_str)
        # [A2] Warn on column mismatch BEFORE pad/trim so the model sees
        # whether row width really matched --headers.
        if headers:
            _warn_col_mismatch_docx(
                expected=len(headers), got=len(cells),
                where=f"row[{ri + 1}]",
            )
            if len(cells) < len(headers):
                cells = cells + [""] * (len(headers) - len(cells))
            elif len(cells) > len(headers):
                cells = cells[: len(headers)]
        cells = [_clip_str(c, _LIM["cell"],
                           where=f"row[{ri}].col[{ci}]")
                 for ci, c in enumerate(cells)]
        # [A3] Consecutive-duplicate detection.
        sig = _row_signature_docx(cells)
        if raw_signatures and sig == raw_signatures[-1] and any(sig):
            _warn_duplicate_row_docx(
                prev_idx=ri, cur_idx=ri + 1, signature=sig,
            )
        raw_signatures.append(sig)
        rows.append(cells)
    if not headers:
        print("[WARN] add-table: --headers is empty; nothing added")
        return
    nc = len(headers)
    tbl = doc.add_table(rows=1, cols=nc)
    # Try a polished built-in style; fall back to plain grid if absent.
    for sty in ("Light Grid Accent 1", "Table Grid"):
        try:
            tbl.style = sty
            break
        except KeyError:
            continue
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        rp = cell.paragraphs[0]
        run = rp.add_run(str(h))
        run.bold = True
    for row in rows:
        rcells = tbl.add_row().cells
        for ci in range(nc):
            rcells[ci].text = row[ci] if ci < len(row) else ""
    # [A4] Tell the model the real shape of the table just inserted.
    _commit_block(
        doc, args, "add-table", meta,
        state_extra=(
            f"last_table=cols={nc} body_rows={len(rows)} "
            f"header='{','.join(headers[:4])}{'...' if len(headers) > 4 else ''}'"
        ),
    )


# --- add-image -------------------------------------------------------------
def cmd_add_image(args):
    """Append an image. Missing file -> placeholder paragraph + [WARN]."""
    doc = _open_doc(args.file)
    meta = _meta_load(doc)
    path = args.path or ""
    if not path or not os.path.exists(path):
        print(f"[WARN] add-image: image not found at '{path}'; "
              "inserting placeholder paragraph instead")
        doc.add_paragraph(f"[image not found: {path}]")
    else:
        try:
            w_in = float(args.width_in if args.width_in is not None else 5.0)
        except (TypeError, ValueError):
            w_in = 5.0
        # Clamp to a sane range so a stray --width-in 9999 doesn't blow up.
        w_in = max(0.5, min(10.0, w_in))
        doc.add_picture(path, width=Inches(w_in))
    _commit_block(doc, args, "add-image", meta)


# --- add-page-break --------------------------------------------------------
def cmd_add_page_break(args):
    """Append a page break."""
    doc = _open_doc(args.file)
    meta = _meta_load(doc)
    doc.add_page_break()
    _commit_block(doc, args, "add-page-break", meta)


# --- catalog ---------------------------------------------------------------
def cmd_catalog(args):
    """Print preset list + add-* cookbook (run ONCE per task)."""
    presets_desc = {
        "business":  "Calibri + deep-navy heading (corporate, US Letter)",
        "academic":  "Times New Roman 12pt (papers, US Letter)",
        "report":    "Arial + red heading (formal report, US Letter)",
        "cjk":       "Microsoft YaHei A4 (Chinese, default)",
    }
    L = []
    L.append("================================================================")
    L.append(" docx_helper.py catalog -- copy-pastable command cookbook       ")
    L.append("================================================================")
    L.append("")
    # COMMON MISTAKES block: shown FIRST so models see the most-frequent
    # argument-format pitfalls before the canonical examples below. The
    # parser also auto-recovers from these via positional fallback +
    # parse_known_args, but warnings still cost a [WARN] line, so prefer
    # the canonical form.
    L.append("=== COMMON MISTAKES (parser auto-recovers, but prefer canonical) ===")
    L.append("  WRONG: add-heading FILE \"Section\" 1 cjk")
    L.append("  RIGHT: add-heading FILE --text \"Section\" --level 1")
    L.append("         (positional fallback also works: add-heading FILE \"Section\" 1)")
    L.append("  WRONG: add-list FILE items a b c")
    L.append("  RIGHT: add-list FILE --type bullet --item a --item b --item c")
    L.append("         (positional fallback also works: add-list FILE a b c)")
    L.append("  WRONG: add-list FILE --type unordered  (or 'ol', 'numbered')")
    L.append("  RIGHT: --type bullet  OR  --type number")
    L.append("         (aliases unordered/ul -> bullet, ordered/ol/numbered -> number)")
    L.append("  WRONG: add-table FILE --rows \"a,b|c,d\"  (--rows alias is tolerated,")
    L.append("                                            but use canonical:)")
    L.append("  RIGHT: add-table FILE --headers \"a,b\" --row \"1,2\" --row \"3,4\"")
    L.append("  WRONG: add-heading FILE --level 1 --text \"X\" --preset cjk")
    L.append("         (--preset is ONLY for new-doc; ignored elsewhere with [WARN])")
    L.append("")
    L.append("=== STEP 1: pick a preset (locks page setup + body/heading "
             "fonts) ===")
    for name in PRESETS:
        L.append(f"    {name:<10}  {presets_desc.get(name, '')}")
    L.append("")
    L.append("=== STEP 2: init an empty doc (creates FILE.docx) ===")
    L.append("  python docx_helper.py new-doc FILE.docx --preset cjk")
    L.append("")
    L.append("=== STEP 3: append blocks (one command per block) ===")
    L.append("")
    L.append("  --- add-title -------------------------------------------------")
    L.append("  Top-of-doc title + optional subtitle. Use ONCE at the start.")
    L.append("    python docx_helper.py add-title FILE.docx \\")
    L.append('      --text     "Q1 2026 Flash Price Report" \\')
    L.append('      --subtitle "Weekly snapshot, 2026-04-27"')
    L.append("")
    L.append("  --- add-heading -----------------------------------------------")
    L.append("  Section heading at level 1..9.")
    L.append("    python docx_helper.py add-heading FILE.docx --level 1 \\")
    L.append('      --text "Executive Summary"')
    L.append("")
    L.append("  --- add-paragraph ---------------------------------------------")
    L.append("  One paragraph. Optional --style 'Quote' / --align "
             "left|center|right|justify.")
    L.append("    python docx_helper.py add-paragraph FILE.docx \\")
    L.append('      --text "Q1 2026 marked our strongest quarter with 18%% '
             'YoY growth."')
    L.append("")
    L.append("  --- add-list --------------------------------------------------")
    L.append("  Bullet (default) or numbered list (--type bullet|number).")
    L.append(f"  1..{_LIM['list_items_max']} items, each <={_LIM['list_item']}"
             " chars.")
    L.append("    python docx_helper.py add-list FILE.docx --type bullet \\")
    L.append('      --item "NAND prices flat WoW" \\')
    L.append('      --item "DRAM contract +1.2%%" \\')
    L.append('      --item "QLC 1Tb stable at $27"')
    L.append("")
    L.append("  --- add-table -------------------------------------------------")
    L.append("  Headed table. Headers/rows are comma-separated.")
    L.append(f"  Max {_LIM['headers_max']} cols x {_LIM['rows_max']} rows; "
             "longer auto-truncated with [WARN].")
    L.append("    python docx_helper.py add-table FILE.docx \\")
    L.append('      --headers "Product,Spot,WoW" \\')
    L.append('      --row     "1Tb QLC,$27.00,0%%" \\')
    L.append('      --row     "1Tb TLC,$29.00,+0.5%%"')
    L.append("")
    L.append("  --- add-image -------------------------------------------------")
    L.append("  Insert an image; --width-in default 5.0 (clamped to 0.5..10).")
    L.append("    python docx_helper.py add-image FILE.docx \\")
    L.append('      --path /sdcard/chart.png --width-in 5')
    L.append("")
    L.append("  --- add-page-break --------------------------------------------")
    L.append("  Insert a page break.")
    L.append("    python docx_helper.py add-page-break FILE.docx")
    L.append("")
    L.append("=== STEP 4: optional -- inspect or edit the result ===")
    L.append("  python docx_helper.py inspect       FILE.docx")
    L.append("  python docx_helper.py extract       FILE.docx")
    L.append("  # set-paragraph FILE IDX TEXT     (positional)")
    L.append("  python docx_helper.py set-paragraph FILE.docx 3 \"New text\"")
    L.append("  # set-cell FILE TABLE ROW COL TEXT  (positional)")
    L.append("  python docx_helper.py set-cell      FILE.docx 0 1 2 \"new cell\"")
    L.append("")
    L.append("=== Hard caps (auto-truncate, never fail) ===")
    L.append(f"  title <={_LIM['title']}    subtitle <={_LIM['subtitle']}    "
             f"heading <={_LIM['heading']}")
    L.append(f"  paragraph <={_LIM['paragraph']}    list item <="
             f"{_LIM['list_item']} x {_LIM['list_items_max']} max")
    L.append(f"  table {_LIM['headers_max']} cols x {_LIM['rows_max']} rows, "
             f"cell <={_LIM['cell']}")
    # %% is escape for the prints above (Python doesn't auto-format here, but
    # leaving %% so copy-paste into shell-style prompts won't be misread).
    print("\n".join(L).replace("%%", "%"))
    _next_step([
        "Now run new-doc, then append one block per add-* command:",
        "  python docx_helper.py new-doc ${WORKSPACE}/doc.docx --preset cjk",
        "  python docx_helper.py add-title ${WORKSPACE}/doc.docx "
        "--text \"...\" --subtitle \"...\"",
        "  python docx_helper.py add-heading / add-paragraph / add-list / "
        "add-table / ...",
    ])




# ===========================================================================
# Reference manual the model reads after every command. We just dump the
# sibling SKILL.md (frontmatter + first H1 stripped) so there is a SINGLE
# source of truth for the CLI surface; updating SKILL.md updates every
# command's tail-output.
# ===========================================================================
_SKILL_MD_BODY_CACHE = None


def _load_skill_md_body():
    """Read sibling SKILL.md, strip YAML frontmatter and the first H1 title.
    Cached for the lifetime of the process. Returns "" on any IO error."""
    global _SKILL_MD_BODY_CACHE
    if _SKILL_MD_BODY_CACHE is not None:
        return _SKILL_MD_BODY_CACHE
    body = ""
    try:
        here = os.path.dirname(os.path.abspath(__file__))
        md_path = os.path.normpath(os.path.join(here, os.pardir, "SKILL.md"))
        with open(md_path, "r", encoding="utf-8") as f:
            text = f.read()
        lines = text.splitlines()
        i, n = 0, len(lines)
        # 1) skip YAML frontmatter delimited by leading `---`
        if i < n and lines[i].strip() == "---":
            i += 1
            while i < n and lines[i].strip() != "---":
                i += 1
            if i < n:
                i += 1
        # 2) skip blank lines after frontmatter
        while i < n and not lines[i].strip():
            i += 1
        # 3) skip the first H1 title line (e.g. "# Docx-Editor")
        if i < n and lines[i].lstrip().startswith("# "):
            i += 1
        # 4) skip blanks after the title
        while i < n and not lines[i].strip():
            i += 1
        body = "\n".join(lines[i:]).rstrip() + "\n"
    except Exception:
        body = ""
    _SKILL_MD_BODY_CACHE = body
    return body


def _print_commands_menu(stream=None):
    """Print SKILL.md (without frontmatter/title) so the model always has the
    full skill manual at the tail of every command's output. Called on both
    success and failure paths."""
    if stream is None:
        stream = sys.stdout
    body = _load_skill_md_body()
    stream.write("\n=== SKILL.md ===\n")
    if body:
        stream.write(body)
    else:
        stream.write("(SKILL.md not found beside this script; "
                     "open skills/docx-editor/SKILL.md for the manual)\n")


# Hallucinated subcommand -> intent-aware "did you mean" hint. Argparse's
# default invalid-choice error already lists every legal cmd, but the model
# still picks the wrong one when its naming intuition does not match.
_DOCX_DID_YOU_MEAN = {
    "new":              "new-doc FILE.docx --preset cjk",
    "new-docx":         "new-doc FILE.docx --preset cjk",
    "new-document":     "new-doc FILE.docx --preset cjk",
    "new-word":         "new-doc FILE.docx --preset cjk",
    "create-doc":       "new-doc FILE.docx --preset cjk",
    "add-text":         "add-paragraph FILE.docx \"text...\"  "
                        "(or  add-heading FILE.docx 1 \"title\")",
    "set-text":         "set-paragraph FILE.docx IDX \"new text\"  "
                        "(run inspect FILE first to find IDX)",
    "write-paragraph":  "add-paragraph FILE.docx \"text...\"",
    "write-table":      "add-table FILE.docx \"a,b\" \"1,2\" \"3,4\"",
    "create-table":     "add-table FILE.docx \"a,b\" \"1,2\" \"3,4\"",
}


def _intercept_hallucinated_cmd(known_cmds):
    """If sys.argv[1] is a well-known wrong name, print a directed fix and
    exit(2) BEFORE argparse runs. Falls through silently otherwise."""
    if len(sys.argv) < 2:
        return
    cmd = sys.argv[1]
    if cmd in known_cmds or cmd.startswith("-"):
        return
    if cmd in _DOCX_DID_YOU_MEAN:
        sys.stderr.write(
            f"[ERROR] unknown subcommand '{cmd}'. "
            f"Did you mean:  {_DOCX_DID_YOU_MEAN[cmd]}\n"
        )
        _print_commands_menu(stream=sys.stderr)
        sys.exit(2)


class SmartArgumentParser(argparse.ArgumentParser):
    def error(self, message):
        sys.stderr.write(f"[ERROR] {message}\n")
        # Print the full commands menu so the model can pick a different
        # subcommand (or fix args of the same one).
        _print_commands_menu(stream=sys.stderr)
        sys.exit(2)


# ===========================================================================
# Argparse wiring
# ===========================================================================
_DOCX_KNOWN_CMDS = {
    "catalog", "new-doc",
    "add-title", "add-heading", "add-paragraph", "add-list",
    "add-table", "add-image", "add-page-break",
    "inspect", "extract", "set-paragraph", "set-cell",
    "new-from-outline", "examples", "create",
}


def main():
    # Intercept well-known wrong subcommand names BEFORE argparse so the model
    # gets a directed fix instead of a generic "invalid choice" listing.
    _intercept_hallucinated_cmd(_DOCX_KNOWN_CMDS)

    ap = SmartArgumentParser(prog="docx_helper",
                             description="Programmatic helper for .docx files.")
    sp = ap.add_subparsers(dest="cmd", required=True)

    # --- new path: catalog + new-doc + add-* (one block per command) -------
    p_cat = sp.add_parser("catalog",
                          help="Cookbook: presets + add-* commands "
                               "(run ONCE per task)")
    p_cat.set_defaults(func=cmd_catalog)

    p_nd = sp.add_parser("new-doc",
                         help="Init an empty docx with a locked preset")
    p_nd.add_argument("file", help="output .docx path")
    p_nd.add_argument("--preset", default="cjk",
                      help=f"one of: {list(PRESETS)} (default cjk)")
    p_nd.set_defaults(func=cmd_new_doc)

    # NOTE on LLM-friendliness: every add-* command below treats its keyword
    # args (--text / --headers / --item / --row) as OPTIONAL at the argparse
    # level and provides a positional fallback. The cmd_* function still
    # validates the resolved value is non-empty, so:
    #   add-heading FILE.docx "标题" 1     == add-heading FILE.docx --text "标题" --level 1
    #   add-list    FILE.docx a b c       == add-list FILE.docx --item a --item b --item c
    #   add-table   FILE.docx "a,b" "1,2" == add-table FILE.docx --headers "a,b" --row "1,2"
    # Unknown arguments (e.g. spurious --preset on add-heading) are tolerated
    # via parse_known_args + a [WARN] line, never an exit-code-2 failure.
    p_at = sp.add_parser("add-title",
                         help="Append a centered title (+ optional subtitle)")
    p_at.add_argument("file")
    p_at.add_argument("pos_text", nargs="?", default=None,
                      help="title text (positional fallback for --text)")
    p_at.add_argument("pos_subtitle", nargs="?", default=None,
                      help="subtitle (positional fallback for --subtitle)")
    p_at.add_argument("--text", default=None)
    p_at.add_argument("--subtitle", default=None)
    p_at.set_defaults(func=cmd_add_title)

    p_ah = sp.add_parser("add-heading",
                         help="Append a heading at level 1..9")
    p_ah.add_argument("file")
    # Two optional positional slots; we sniff which is the int level vs text
    # in cmd_add_heading so both 'FILE TEXT LEVEL' and 'FILE LEVEL TEXT' work.
    p_ah.add_argument("pos1", nargs="?", default=None,
                      help="text or level (positional fallback)")
    p_ah.add_argument("pos2", nargs="?", default=None,
                      help="text or level (positional fallback)")
    p_ah.add_argument("--level", type=int, default=None)
    p_ah.add_argument("--text", default=None)
    p_ah.set_defaults(func=cmd_add_heading)

    p_ap = sp.add_parser("add-paragraph",
                         help="Append a paragraph (optional --style/--align)")
    p_ap.add_argument("file")
    p_ap.add_argument("pos_text", nargs="?", default=None,
                      help="paragraph text (positional fallback for --text)")
    p_ap.add_argument("--text", default=None)
    p_ap.add_argument("--style", default=None,
                      help="paragraph style: 'Quote' / 'Intense Quote' / "
                           "'Normal' / 'List Bullet' / etc.")
    p_ap.add_argument("--align", default=None,
                      choices=["left", "center", "right", "justify"],
                      help="paragraph alignment")
    p_ap.set_defaults(func=cmd_add_paragraph)

    p_al = sp.add_parser("add-list",
                         help="Append bullet (default) or numbered list "
                              "(1..50 items)")
    p_al.add_argument("file")
    # Free-form positional items; merged with --item in cmd_add_list.
    p_al.add_argument("pos_items", nargs="*",
                      help="list items (positional fallback for --item)")
    # No 'choices=' constraint -- cmd_add_list normalizes synonyms
    # (unordered/ul/disc -> bullet, ordered/ol/numbered -> number).
    p_al.add_argument("--type", default="bullet")
    p_al.add_argument("--item", action="append", default=[],
                      help="repeat for each list item")
    p_al.set_defaults(func=cmd_add_list)

    p_atb = sp.add_parser("add-table",
                          help="Append a table from --headers and --row(s)")
    p_atb.add_argument("file")
    # Free-form positional fallback: first slot=headers, rest=rows.
    p_atb.add_argument("pos_args", nargs="*",
                       help="positional fallback: HEADERS [ROW1 ROW2 ...]")
    p_atb.add_argument("--headers", default=None,
                       help="comma-separated header row")
    p_atb.add_argument("--row", action="append", default=[],
                       help="comma-separated body row; repeat per row")
    # --rows alias: model frequently writes --rows "a,b|c,d"; we split by '|'.
    p_atb.add_argument("--rows", action="append", default=[],
                       help="alias for --row; '|' separates multiple rows")
    p_atb.set_defaults(func=cmd_add_table)

    p_aim = sp.add_parser("add-image", help="Append an image")
    p_aim.add_argument("file")
    p_aim.add_argument("--path", required=True)
    p_aim.add_argument("--width-in", dest="width_in", type=float, default=5.0)
    p_aim.set_defaults(func=cmd_add_image)

    p_apb = sp.add_parser("add-page-break", help="Append a page break")
    p_apb.add_argument("file")
    p_apb.set_defaults(func=cmd_add_page_break)

    # --- existing commands (kept for backward compat) ----------------------
    p_ins = sp.add_parser("inspect", help="Per-paragraph + per-table listing")
    p_ins.add_argument("file"); p_ins.set_defaults(func=cmd_inspect)

    p_ext = sp.add_parser("extract", help="Plain text dump")
    p_ext.add_argument("file"); p_ext.set_defaults(func=cmd_extract)

    # set-paragraph FILE IDX TEXT  (canonical positional form)
    p_st = sp.add_parser("set-paragraph",
                         help="Replace paragraph text: FILE IDX TEXT")
    p_st.add_argument("file")
    p_st.add_argument("idx_pos", nargs="?", default=None,
                      help="paragraph index from `inspect`")
    p_st.add_argument("text_pos", nargs="?", default=None,
                      help="new text; '\\n' produces a soft line break")
    p_st.add_argument("--style", default=None,
                      help="optional style: 'Heading 1', 'Quote', etc.")
    # Legacy flags (hidden; still functional).
    p_st.add_argument("--idx", type=int, default=None, help=argparse.SUPPRESS)
    p_st.add_argument("--text", default=None, help=argparse.SUPPRESS)
    p_st.add_argument("--out", "--output", dest="out", default=None,
                      help=argparse.SUPPRESS)
    p_st.set_defaults(func=cmd_set_paragraph)

    # set-cell FILE TABLE ROW COL TEXT  (canonical positional form)
    p_sc = sp.add_parser("set-cell",
                         help="Replace table cell text: FILE TABLE ROW COL TEXT")
    p_sc.add_argument("file")
    p_sc.add_argument("table_pos", nargs="?", default=None,
                      help="table index (0-based) from `inspect`")
    p_sc.add_argument("row_pos", nargs="?", default=None,
                      help="row index (0-based)")
    p_sc.add_argument("col_pos", nargs="?", default=None,
                      help="column index (0-based)")
    p_sc.add_argument("text_pos", nargs="?", default=None,
                      help="new cell text")
    # Legacy flags (hidden; still functional).
    p_sc.add_argument("--table", type=int, default=None, help=argparse.SUPPRESS)
    p_sc.add_argument("--row", type=int, default=None, help=argparse.SUPPRESS)
    p_sc.add_argument("--col", type=int, default=None, help=argparse.SUPPRESS)
    p_sc.add_argument("--text", default=None, help=argparse.SUPPRESS)
    p_sc.add_argument("--out", "--output", dest="out", default=None,
                      help=argparse.SUPPRESS)
    p_sc.set_defaults(func=cmd_set_cell)

    p_nw = sp.add_parser("new-from-outline",
                         help="Build a docx from an outline JSON")
    p_nw.add_argument("outline_pos", nargs="?", default=None,
                      help="outline JSON path (positional)")
    p_nw.add_argument("--outline", "--input", dest="outline_opt", default=None,
                      help="alias for the positional outline path")
    p_nw.add_argument("--out", "--output", dest="out", default=None,
                      help="output .docx (omit when --dry-run)")
    p_nw.add_argument("--preset", default=None,
                      help=f"one of: {list(PRESETS)} (default 'business')")
    p_nw.add_argument("--dry-run", action="store_true",
                      help="validate outline only; no file written")
    p_nw.set_defaults(func=cmd_new_from_outline)

    p_ex = sp.add_parser("examples",
                         help="List or dump built-in outline JSON examples")
    p_ex.add_argument("name", nargs="?", default="list",
                      help=f"one of: {list(EXAMPLES_DOCX)} or 'list'")
    p_ex.add_argument("--out", "--output", dest="out", default=None,
                      help="write example outline JSON to this file "
                           "(default: print to stdout)")
    p_ex.set_defaults(func=cmd_examples)

    p_cr = sp.add_parser("create",
                         help="Quick path: generate a doc from a template in one shot")
    p_cr.add_argument("--template", required=True,
                      help=f"one of: {list(EXAMPLES_DOCX)}")
    p_cr.add_argument("--out", "--output", dest="out", required=True,
                      help="output .docx path")
    p_cr.add_argument("--outline", "--input", dest="outline_file",
                      default=None,
                      help="optional outline JSON that overrides the template")
    p_cr.add_argument("--preset", default=None)
    p_cr.set_defaults(func=cmd_create)

    # Use parse_known_args so unknown flags (e.g. spurious --preset on
    # add-heading, which models hallucinate from the new-doc example) are
    # tolerated with a [WARN] line instead of exit-code-2 argparse failure.
    args, unknown = ap.parse_known_args()
    if unknown:
        print(f"[WARN] ignoring unknown args for {args.cmd}: {unknown}")
    try:
        args.func(args)
        # Print the commands menu so the model can pick the next call.
        # We do NOT presume to suggest a specific next step -- the helper
        # has no idea what the overall task is.
        _print_commands_menu()
    except Exception as e:
        sys.stderr.write(f"[ERROR] {type(e).__name__}: {e}\n")
        _print_commands_menu(stream=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
